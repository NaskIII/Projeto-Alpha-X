import docx
from docx.shared import Length


class Docx(object):

    def docx(self, lista, caminho, title):  # Método que irá criar um .docx
        documento = docx.Document()  # Chamo a classe
        style = documento.styles['Normal']  # Defino o estilo para normal
        font = style.font  # Chamo a classe fonte
        font.name = 'Arial'  # Mudo a fonte para Arial
        font.size = docx.shared.Pt(13)  # Mudo o tamanho da fonte
        paragrafo = documento.add_paragraph()
        formato_paragrafo = paragrafo.paragraph_format
        formato_paragrafo.line_spacing = 2

        documento.add_heading(title, 0)  # Adiciono um título principal na primeira página

        for i in lista:  # Itero a lista para escrever no docx
            if i.__contains__('==') and len(i) <= 7:  # Se o conteúdo do índice da lista, for menor que 7, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=0)
            elif i.__contains__('==') and len(i) <= 10:  # Se o conteúdo do índice da lista, for menor que 10, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=1)
            elif i.__contains__('==') and len(i) <= 20:  # Se o conteúdo do índice da lista, for menor que 20, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=2)
            elif i.__contains__('==') and len(i) <= 35:  # Se o conteúdo do índice da lista, for menor que 35, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=3)
            elif i.__contains__('==') and len(i) > 35:
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=3)
            else:  # Se nenhuma for verdadeira, adiciona apenas um parágrafo
                documento.add_paragraph(i)

        documento.save(caminho + '/' + title + '.docx')  # Salvo o documento

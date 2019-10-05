'''
Autor: Raphael Nascimento
ID: Nask
Objetivo: Buscar termos na Wikipedia e retornar os resultados em formato .txt, alem de
gerar um resumo do conteudo original. No final serao gerados 2 arquivos, um contendo o conteudo do Wikipedia
e o outro o resumo.

FLUXO DE CHAMADAS DOS METODOS!!!
    1 - self.resumir()
    2 - self.formatarTexto()
    3 - countLines()
    4 - self.write()
    5 - self.wiki()
    6 - self.cleanSentences()
    7 - formatarReferencias()
    
    FLUXO DE FUNCIONAMENTO DO CODIGO!!!
    1 - self.wiki()
    2 - self.write()
    3 - self.cleanSentences()
    4 - contLines()
    5 - self.formatarTexto()
    6 - self.formatarReferencias()
    7 - self.resumir()
'''

import sys
import Algorithmia  # API usada para buscar e resumir o conteudo da wikipedia
from robots import Diretorios
from robots import Wikipedia
import docx
import os
from docx.shared import Length


class TextRobots(object):  # Classe responsavel por gerar todo o conteudo de texto
    def __init__(self, artigo):
        self.artigo = artigo
        self.caminho = None
        self.content = None
        self.lista = []
        self.wikipedia = Wikipedia.Wikipedia(self.artigo)

    def wiki(self):  # Metodo usado para pegar o conteudo
        content = self.wikipedia.content()
        try:
            dire = Diretorios.start(self.artigo['termo'])
            self.caminho = dire
            return content
        except:
            print('O conteudo informado nao foi encontrado, por favor busque com outras palavras.')
            sys.exit()

    def cleanSentences(self, texto):  # Retiro as sentencas de marcaçao utilizado no Wikipedia
        texto = texto.replace('Referências', '')
        return texto

    def write(self, content):  # Escrevo o conteudo em um arquivo .txt para poder formatar linha por linha
        self.content = content
        texto = self.cleanSentences(self.content)
        new_arq = open(self.caminho + self.wikipedia.title() + '.txt', 'w', encoding='utf-8')
        new_arq.writelines(texto)
        new_arq.close()
        return content
        

    def contLines(self):  # Itero o arquivo para poder saber o numero de linhas
        arq = open(self.caminho + self.wikipedia.title() + '.txt', 'r', encoding='utf-8')
        linhas = len(arq.readlines())
        arq.close()
        return linhas

    def formatarReferencias(self):  # Retiro as marcaçoes que estao nas referencias
        referencias = str(self.wikipedia.references())
        referencias = referencias.replace('[', '')
        referencias = referencias.replace(']', '')
        referencias = referencias.replace("'", '')
        referencias = referencias.replace(',', '\n')
        return referencias

    def formatarTexto(self):  # Formato o texto para ser reescrito no .txt
        linhas = self.contLines()
        texto_final = ""
        arq = open(self.caminho + self.wikipedia.title() + '.txt', 'r', encoding='utf-8')

        for i in range(linhas):
            texto = str(arq.readline())
            if texto != '\n':
                texto = texto.lstrip(" ")
                texto_final += texto

        texto_final.replace('Referências', '')
        arq.close()
        new_arq = open(self.caminho + self.wikipedia.title() + '.txt', 'w', encoding='utf-8')
        new_arq.writelines(texto_final)
        new_arq.writelines('\n')
        new_arq.writelines('\n')
        new_arq.writelines('Referências:\n')
        new_arq.writelines(self.formatarReferencias())
        new_arq.close()
        texto_final = texto_final.replace('\n', '\n\n')
        return texto_final

    def resumir(self, texto):  # Metodo utilizado para resumir uma string

        input = texto, 50  # Variavel utilizada para dar entrada no algoritmo
        client = Algorithmia.client('simyw+zYbXC1hUyLm4AVdUorUMD1')  # Faço minha autenticaçao na API
        algo = client.algo('nlp/Summarizer/0.1.8')  # Chamo o Algoritmo que faz resumos
        algo.set_options(timeout=300)  # Defino um tempo maximo de resposta, opcional
        resumo = algo.pipe(input).result  # Recupero o texto resumido

        arquivo = open(self.caminho + self.wikipedia.title() + '_Resumido.txt',
                       'w', encoding='utf-8')  # Gero um novo arquivo para o novo conteudo
        arquivo.writelines(resumo)  # Escrevo o mesmo em um novo arquivo
        arquivo.close()  # Fecha o Arquivo

    def chamadas(self):  # Metodo que vai chamar os outros metodos da classe
        content = self.wiki()
        self.write(content)
        texto = self.formatarTexto()
        self.resumir(texto)
        self.read()
        self.docx()

    def read(self):  # Irá ler o txt novamente e colocar o conteudo dentro de uma lista
        arquivo = open(self.caminho + self.wikipedia.title() + '.txt', 'r', encoding='utf-8')
        linhas = self.contLines()

        for i in range(linhas):
            self.lista.append(str(arquivo.readline()))
        arquivo.close()

    def docx(self):  # Método que irá criar um .docx 
        documento = docx.Document()  # Chamo a classe
        style = documento.styles['Normal']  # Defino o estilo para normal
        font = style.font  # Chamo a classe fonte
        font.name = 'Arial'  # Mudo a fonte para Arial
        font.size = docx.shared.Pt(13)  # Mudo o tamanho da fonte
        paragrafo = documento.add_paragraph()
        formato_paragrafo = paragrafo.paragraph_format
        formato_paragrafo.line_spacing = 2

        documento.add_heading(self.wikipedia.title(), 0)  # Adiciono um título principal na primeira página

        for i in self.lista:  # Itero a lista para escrever no docx
            if i.__contains__('==') and len(i) <= 7:  # Se o conteúdo do índice da lista, for menor que 7, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=0)
            elif i.__contains__('==') and len(i) <= 10:  # Se o conteúdo do índice da lista, for menor que 10, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=1)  
            elif i.__contains__('==') and len(i) <= 20:  # Se o conteúdo do índice da lista, for menor que 20, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=2)
            elif i.__contains__('==') and len(i) <= 35:  # Se o conteúdo do índice da lista, for menor que 35, adiciona um título maior
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=3)
            elif i.__contains__('==') and len(i) > 35:
                i = i.replace('=', '')
                i = i.lstrip(' ')
                documento.add_heading(i, level=3)
            else:  # Se nenhuma for verdadeira, adiciona apenas um parágrafo 
                documento.add_paragraph(i)

        documento.save(self.caminho + '/' + self.wikipedia.title() + '.docx')  # Salvo o documento
        self.apagar() 


    def apagar(self):  # Apago o txt com o mesmo conteúdo do docx
        os.remove(self.caminho + self.wikipedia.title() + '.txt')

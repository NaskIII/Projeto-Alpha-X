'''
Autor: Raphael Nascimento
ID: Nask
Objetivo: Pegar o Texto formatado de um .txt e reescreve-lo no formato docx
'''


import docx


class Docx(object):
    def __init__(self, caminho, diretorio, content):
        self.caminho = caminho
        self.diretorio = diretorio
        self.content = content
        self.lista = []

    def contLines(self):  # Itero o arquivo para poder saber o numero de linhas
        arq = open(self.caminho , 'r')
        linhas = len(arq.readlines())
        arq.close()
        return linhas

    def read(self):
        arquivo = open(self.caminho, 'r')
        linhas = self.contLines()

        for i in range(linhas):
            self.lista.append(str(arquivo.readline()))
        arquivo.close()

    def docx(self):
        documento = docx.Document()
        style = documento.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)

        documento.add_heading(self.content['title'], 0)

        documento.add_paragraph(self.lista[0])

        for i in self.lista:
            if 7 >= len(i) > 0:
                documento.add_heading(i, level=0)
            elif 10 >= len(i) > 0:
                documento.add_heading(i, level=1)
            elif 20 >= len(i) > 0:
                documento.add_heading(i, level=2)
            elif 30 >= len(i) > 0:
                documento.add_heading(i, level=3)
            else:
                documento.add_paragraph(i)
        documento.save(self.diretorio + '/' + self.content['title'] + '.docx')

    def chamadas(self):
        self.read()
        self.docx()

